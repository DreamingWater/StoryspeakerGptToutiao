from io import StringIO

from docx import Document
from docx.shared import Inches
import os
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


# https://python-docx.readthedocs.io/en/latest/index.html

def docx_name_deal(name):
    if name.endswith('docx') == False and name.endswith('doc') == False:
        name += '.docx'
    return name


class WrodCMD:
    def __init__(self):
        self.document = Document()


class WriteWord(WrodCMD):
    def __init__(self):
        super().__init__()
        self.output_dir = os.path.split(os.path.realpath(__file__))[0] + '\\Output\\'

    def add_heading(self, heading):
        self.document.add_heading(heading, 1)

    def add_paragragh(self, paragragh):
        style = self.document.styles['Normal']
        # 设置西文字体
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        # 设置首行缩进， 先获取段落样式
        paragraph_format = style.paragraph_format
        # 首行缩进0.74厘米，即2个字符
        paragraph_format.first_line_indent = Cm(0.74)
        paragraph_format.space_after = Pt(0)  # 段后间距
        paragraph_format.space_before = Pt(5)  # 段后间距
        self.document.add_paragraph(paragragh)

    def save_document(self, name: str):
        name = docx_name_deal(name)
        self.document.save(self.output_dir + name)


class ReadWord(WrodCMD):
    def __init__(self):
        super().__init__()
        self.chapter = {}
        self.paragraphs = None # docx里面的内容
    def open_word(self, file_name):
        file_name = docx_name_deal(file_name)
        self.document = Document(file_name)
        self.paragraphs = self.document.paragraphs
    def get_sections(self):  # 获取文档所有number of sections
        sections = self.document.sections
        print("the number of sections is {}".format(len(sections)))

    def get_paragraphts(self,index):
        return self.paragraphs[index].text
    def get_all_paragraphs(self):

        print("the number of paragraphs is {}".format(len(self.paragraphs)))  # 打印结果：20
        pars_string = [par.text for par in self.paragraphs]
        return ''.join(pars_string)
        # for index, par in enumerate(pars_string):
            #print(index, par)

    # 区分出head和文本段落，并联合组合为chapter，保存到self.chapter
    def distinguish_head_para(self):
        paragraphs = self.document.paragraphs
        paragraphs_number = len(paragraphs)  # get the 92number of paragraph
        this_chapter_border = []  # 此列表用于保存head的index
        for index, paragraph in enumerate(paragraphs):
            style_name = paragraph.style.name
            if style_name.startswith('Heading'):  # 判断是不是head
                this_chapter_border.append(index)
        if 0 not in this_chapter_border:  # 保证从第一段开始都进行分章节
            this_chapter_border.insert(0, 0)
        # check the chapter
        this_chapter_border.append(paragraphs_number)  # 在此处将最后一段的index加入到，后续就可以直接切割
        for index in range(len(this_chapter_border) - 1):
            this_range = range(this_chapter_border[index], this_chapter_border[index + 1])  # 字典的value包括key的内容
            range_list = [i for i in this_range]
            if len(this_range) <= 1:
                print('There if an error in the chapter head..')
                print('please check the {} paragragh...'.format(index))
            self.chapter[this_chapter_border[index]] = range_list
        print('the chapter of this is as following-----', end='    ')
        print(self.chapter)  #{'0':[0,1],'2':[2]}
        # return self.chapter

    def get_all_question(self):
        questions = []
        for key in self.chapter.keys():
            if key == 0:
                back_ground = self.chapter.get(key)
                paras = [self.get_paragraphts(i) for i in back_ground]
                questions.append(''.join(paras))
            else:
                questions.append(self.get_paragraphts(key))
        return questions
